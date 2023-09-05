# import docx
from docx import Document
import re
import os
import docx
from docx.enum.style import WD_STYLE_TYPE
# import PyPDF2
# import textract
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
# from win32com.client import gencache, constants
import xml.etree.ElementTree as ET
import docModel
# import figModel
import docProcData
import win32com
from win32com.client import Dispatch
import threading
from gensim import corpora
from gensim import models
from gensim import similarities
import jieba
from docx.enum.table import WD_TABLE_ALIGNMENT

os.environ['TF_CPP_MIN_LOG_LEVEL'] = '2'


def start_word(path):
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    doc = word.Documents.Open(FileName=path, Encoding='gbk')
    return word, doc


def close_word(doc, word):
    doc.Close()
    word.Quit()
    print("close word")


def add_comment(doc, position1, comment):
    # doc = word.Documents.Open(FileName=path, Encoding='gbk')
    # 主要关键的是这一句
    '''find_t = position1
    word.Selection.Find.Execute(find_t)
    doc.Comments.Add(Range=word.Selection.Range, Text=comment)'''
    for each in doc.paragraphs:
        print(each.Range.text)
        if position1 in each.Range.text:
            print("here")
            doc.Comments.Add(Range=each.Range, Text=comment)
            doc.Save()


def find_title(doc_list):
    content = doc_list
    tree_of_word = dict()
    tree_of_title = dict()
    tree_of_title_1 = dict()
    big_title = ""
    small_title = ""
    for each in content:
        if len(each[1]) > 0:
            for letter in each[1]:
                if letter != " ":
                    if each[0] == "Heading 1":
                        big_title = each[1]
                        tree_of_title[big_title] = []
                    elif each[0] == "Heading 2":
                        small_title = each[1]
                        tree_of_title_1[small_title] = big_title
                        if big_title != "":
                            tree_of_title[big_title].append(small_title)
                    else:
                        tree_of_word[str(each[1])] = big_title + " " + small_title
                    break
    return tree_of_word, tree_of_title, tree_of_title_1


'''def createPdf(wordPath, pdfPath):
    """
    word转pdf
    :param wordPath: word文件路径
    :param pdfPath:  生成pdf文件路径
    """
    word = gencache.EnsureDispatch('Word.Application')
    doc = word.Documents.Open(wordPath, ReadOnly=1)
    doc.ExportAsFixedFormat(pdfPath,
                            constants.wdExportFormatPDF,
                            Item=constants.wdExportDocumentWithMarkup,
                            CreateBookmarks=constants.wdExportCreateHeadingBookmarks)
    word.Quit(constants.wdDoNotSaveChanges)'''

'''def getpage(word_file, keyword):
    pdf_file = os.path.splitext(word_file)[0] + ".pdf"
    createPdf(word_file, pdf_file)
    print("successfully created")
    # Open the pdf file
    object = PyPDF2.PdfFileReader(pdf_file)
    # Get number of pages
    NumPages = object.getNumPages()
    # Extract text and do the search
    for i in range(0, NumPages):
        PageObj = object.getPage(i)
        Text = PageObj.extractText()
        if Text == "":
            Text = textract.process(PageObj, method='tesseract', language='eng')
        if re.search(keyword, Text):
            print("Pattern Found on Page: " + str(i))
    os.remove(pdf_file)'''


def write_paragraph(word_path, doc_result):
    f = open(word_path, 'w', encoding='utf-8')
    content = ''
    if len(doc_result) > 0:
        for each in doc_result:
            content += each + '\n'
    else:
        content += "未找到审查意见。"
    f.write(content)
    f.close()


def get_paragraph(word_path):
    doc = docx.Document(word_path)
    paragraphs = doc.paragraphs
    for paragraph in paragraphs:
        print(paragraph)
    return paragraphs


def get_pictures(word_path, result_path):
    """
    图片提取
    :param word_path: word路径
    :param result_path: 结果路径
    :return:
    """
    doc = docx.Document(word_path)
    dict_rel = doc.part._rels
    for rel in dict_rel:
        rel = dict_rel[rel]
        if "image" in rel.target_ref:
            if not os.path.exists(result_path):
                os.makedirs(result_path)
            img_name = re.findall("/(.*)", rel.target_ref)[0]
            word_name = os.path.splitext(word_path)[0]
            if os.sep in word_name:
                new_name = word_name.split('\\')[-1]
            else:
                new_name = word_name.split('/')[-1]
            # img_name = f'{new_name}_{img_name}'
            img_name = f'{img_name.split("image")[1]}'
            with open(f'{result_path}/{img_name}', "wb") as f:
                f.write(rel.target_part.blob)


def get_form(path):
    document = Document(path)  # 读入文件
    tables = document.tables  # 获取文件中的表格集

    for table in tables[:]:
        for i, row in enumerate(table.rows[:]):  # 读每行
            row_content = []
            for cell in row.cells[:]:  # 读一行中的所有单元格
                c = cell.text
                row_content.append(c)
            print(row_content)  # 以列表形式导出每一行数据


def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def read_table(table):
    return [[cell.text for cell in row.cells] for row in table.rows]


def read_word(word_path):
    get_pictures(word_path, R".\fig_process")
    doc_list = []
    doc = docx.Document(word_path)
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            if "Heading" in block.style.name:
                # print(block.style.name)
                # print("title", [block.text])
                doc_list.append((block.style.name, block.text))
            else:
                # print("text", [block.text])
                if len(block.style.name.split("、")) == 2:
                    doc_list.append(("Heading 1", block.text))
                elif len(block.style.name.split(".")) == 2:
                    doc_list.append(("Heading 2", block.text))
                else:
                    doc_list.append(("text", block.text))
        elif isinstance(block, Table):
            # print("table", read_table(block))
            doc_list.append(("table", read_table(block)))
    return doc_list


def read_xml(path):
    root = ET.parse(path).getroot()
    content_list = []
    for node in root.iter():
        if "body" in node.tag:
            for each in node:
                if each.tag[len(each.tag) - 1] == "p":
                    flag = 0
                    for one in each.iter():
                        if "oMathPara" in one.tag and flag == 0:
                            content_list.append("equation")
                            flag = 1
                        if "drawing" in one.tag and flag == 0:
                            content_list.append("picture")
                            flag = 1
                    if flag == 0:
                        content_list.append("text")
                if each.tag[len(each.tag) - 1] == "l":
                    content_list.append("table")
    # print(content_list)
    return content_list


def check(file_path):
    doc_list = read_word(file_path)
    # get_title(R'C:\Users\Archillesheel\Desktop\umji\ve492\homework\test.docx')
    # content_list = read_xml(R'C:\Users\Han Shaochen\Desktop\sjtu\byy_prj\样本\示例2\长青线路跨越大广高速专项施工方案.xml')
    f = open("docProcParaData.txt", 'w', encoding='utf-8')
    for each in doc_list:
        if each[0] == "text" and len(each[1]) > 0:
            # print(len(each[1]))
            f.write(each[1] + "\n")
    f.close()
    para = docModel.merge_str(docProcData.para_data(0))
    keys_sets = docProcData.keys_data()
    for keys in keys_sets:
        if keys[0] == '文明施工：\n':
            print("number of keys is", len(keys))
            doc_result = docModel.doc_proc(para, keys)
            print("number of unfound keys is", len(doc_result))
            for res in doc_result:
                print(res)
            nok = "number of keys is " + str(len(keys))
            nouk = "number of unfound keys is " + str(len(doc_result))
            return nok, nouk, doc_result


def check_sequence(word, location_list=None):
    try:
        # 获取相似度阈值标准
        f = open("相似度阈值.txt")
        standard = float(f.readline())
        f.close()
    except:
        standard = 0.5
    # 标准库的分词和建模
    all_location_list = []
    content = read_word(word)
    # print(content)
    if location_list == None:
        location_list = []
        for one in content:
            if one[0] == "Heading 2":
                location_list.append(one[1])
        print(location_list)
    for doc in location_list:
        doc_list = [word for word in jieba.cut_for_search(doc)]
        # doc_list = [word for word in jieba.cut(doc)]
        all_location_list.append(doc_list)
    # 1 获取词袋
    dictionary = corpora.Dictionary(all_location_list)
    # 2 制作语料库
    # 标准库的二元组向量转换
    corpus = [dictionary.doc2bow(doc) for doc in all_location_list]
    # 3 使用TF-IDF模型对语料库建模
    tfidf = models.TfidfModel(corpus)
    #####################################################################################
    # 检测文档的分词
    word_location = dict()
    head1 = ""
    head2 = ""
    for each in content:
        if each[0] == "Heading 1" or "、" in each[1] and each[1] != "":
            head1 = each[1]
            word_location[head1] = []
        if each[0] == "Heading 2" and each[1] != "" and head1 != "":
            head2 = each[1]
            if head2 != head1:
                word_location[head1].append(head2)
                word_location[head2] = ""
        if each[0] == "text" and each[1] != "" and head2 != "":
            word_location[head2] += each[1]

    content_index = 0
    wrong_location = 0
    for each in word_location.keys():
        # print(type(word_location[each]))
        if type(word_location[each]) == str:
            doc_test = word_location[each]
            doc_test_list = [word for word in jieba.cut_for_search(doc_test)]
            # doc_test_list = [word for word in jieba.cut(doc_test)]

            # 测试文档的二元组向量转换
            doc_test_vec = dictionary.doc2bow(doc_test_list)

            # 对每个目标文档，分析测试文档的相似度
            index = similarities.SparseMatrixSimilarity(tfidf[corpus], num_features=len(dictionary.keys()))
            sim = index[tfidf[doc_test_vec]]

            # 根3.3 据相似度排序
            sim = sorted(enumerate(sim), key=lambda item: -item[1])
            threshold = [sim[i][0] for i in range(int(len(sim) / 2))]
            if content_index not in threshold:
                for pair in sim:
                    if pair[0] == content_index and pair[1] <= standard:
                        print(each + " wrong location")
                        print("position: ", content_index)
                        wrong_location += 1
                        print(sim)
            content_index += 1
    print(wrong_location)


def out_table(word_path, out_content):
    doc = docx.Document(word_path)
    # table = doc.tables[0]
    # table.alignment = WD_TABLE_ALIGNMENT.CENTER
    out_table = doc.add_table(1, 2)
    out_table.rows[0].cells[0].text = "序号"
    out_table.rows[0].cells[1].text = "意见"
    doc.save(word_path)


if __name__ == '__main__':
    # check_sequence(R"C:\Users\Archillesheel\PycharmProjects\readdoc\长青线路跨越大广高速专项施工方案.docx")
    out_table(R"C:\Users\Archillesheel\Desktop\样本\示例1\test.docx", [])
    # get_paragraph(R"C:\Users\Archillesheel\Desktop\样本\示例1\test.docx")
